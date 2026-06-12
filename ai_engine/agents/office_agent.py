"""
Office Agent - Agente especializado en automatización de Microsoft Office (Word, Excel, PowerPoint).
Permite crear, editar y analizar documentos de Office de forma autónoma.
"""
import asyncio
import logging
import os
from typing import Dict, List, Optional, Any
from pathlib import Path

try:
    from .base_agent import BaseAgent, AgentState, AgentCapability
except ImportError:
    from ai_engine.agents.base_agent import BaseAgent, AgentState, AgentCapability

logger = logging.getLogger(__name__)

class OfficeAgent(BaseAgent):
    """Agente para automatización de tareas de Microsoft Office."""
    
    def __init__(self, agent_id: str = "office-agent"):
        super().__init__(agent_id=agent_id)
        self.name = "OfficeAgent"
        self.description = "Automates Microsoft Office tasks (Word, Excel, PowerPoint)"
        
        # Registrar capacidades
        self.capabilities = [
            AgentCapability.CREATE_DOCUMENT,
            AgentCapability.EDIT_DOCUMENT,
            AgentCapability.READ_DOCUMENT,
            AgentCapability.EXTRACT_DATA,
            AgentCapability.FORMAT_DOCUMENT,
            AgentCapability.CONVERT_FORMAT
        ]
        
        # Estado del agente
        self.current_document: Optional[str] = None
        self.document_type: Optional[str] = None  # 'word', 'excel', 'powerpoint'
        
    async def initialize(self):
        """Inicializa el agente Office."""
        logger.info(f"Initializing {self.name}...")
        
        # Verificar dependencias
        try:
            import comtypes  # Para Windows con COM
            logger.info("COM automation available (Windows)")
        except ImportError:
            try:
                from openpyxl import Workbook  # Alternativa multiplataforma para Excel
                from docx import Document  # Para Word
                logger.info("Using openpyxl/docx libraries (cross-platform)")
            except ImportError:
                logger.warning("Office libraries not installed. Limited functionality.")
                
        self.state = AgentState.IDLE
        logger.info(f"{self.name} initialized successfully.")

    async def execute_task(self, task: Dict[str, Any]) -> Dict[str, Any]:
        """Ejecuta una tarea de Office."""
        if self.state != AgentState.IDLE:
            return {"success": False, "error": "Agent is busy"}
            
        self.state = AgentState.BUSY
        action = task.get("action")
        
        try:
            if action == "create_excel":
                result = await self._create_excel(task)
            elif action == "read_excel":
                result = await self._read_excel(task)
            elif action == "write_excel":
                result = await self._write_excel(task)
            elif action == "create_word":
                result = await self._create_word(task)
            elif action == "read_word":
                result = await self._read_word(task)
            elif action == "edit_word":
                result = await self._edit_word(task)
            else:
                result = {"success": False, "error": f"Unknown action: {action}"}
                
        except Exception as e:
            logger.error(f"Task execution failed: {e}")
            result = {"success": False, "error": str(e)}
        finally:
            self.state = AgentState.IDLE
            
        return result

    async def _create_excel(self, task: Dict) -> Dict:
        """Crea un archivo Excel nuevo."""
        filename = task.get("filename", "report.xlsx")
        data = task.get("data", [])
        headers = task.get("headers", [])
        
        try:
            from openpyxl import Workbook
            
            wb = Workbook()
            ws = wb.active
            
            # Agregar headers
            if headers:
                for col_num, header in enumerate(headers, 1):
                    ws.cell(row=1, column=col_num, value=header)
                    
            # Agregar datos
            for row_num, row_data in enumerate(data, 2 if headers else 1):
                for col_num, value in enumerate(row_data, 1):
                    ws.cell(row=row_num, column=col_num, value=value)
                    
            # Guardar archivo
            filepath = Path(filename)
            if not filepath.is_absolute():
                filepath = Path.cwd() / filepath
                
            wb.save(str(filepath))
            
            logger.info(f"Excel file created: {filepath}")
            return {
                "success": True,
                "filepath": str(filepath),
                "rows": len(data),
                "columns": len(headers) if headers else (len(data[0]) if data else 0)
            }
            
        except ImportError:
            return {"success": False, "error": "openpyxl not installed"}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def _read_excel(self, task: Dict) -> Dict:
        """Lee un archivo Excel."""
        filename = task.get("filename")
        sheet_name = task.get("sheet", 0)
        
        if not filename:
            return {"success": False, "error": "Filename required"}
            
        try:
            from openpyxl import load_workbook
            
            filepath = Path(filename)
            if not filepath.exists():
                return {"success": False, "error": "File not found"}
                
            wb = load_workbook(str(filepath), read_only=True)
            ws = wb[sheet_name] if isinstance(sheet_name, str) else wb.worksheets[sheet_name]
            
            # Leer datos
            data = []
            for row in ws.iter_rows(values_only=True):
                data.append(list(row))
                
            wb.close()
            
            return {
                "success": True,
                "data": data,
                "rows": len(data),
                "columns": len(data[0]) if data else 0
            }
            
        except ImportError:
            return {"success": False, "error": "openpyxl not installed"}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def _write_excel(self, task: Dict) -> Dict:
        """Escribe datos en un Excel existente."""
        filename = task.get("filename")
        data = task.get("data", [])
        start_row = task.get("start_row", 1)
        start_col = task.get("start_col", 1)
        
        if not filename:
            return {"success": False, "error": "Filename required"}
            
        try:
            from openpyxl import load_workbook
            
            filepath = Path(filename)
            if not filepath.exists():
                return {"success": False, "error": "File not found"}
                
            wb = load_workbook(str(filepath))
            ws = wb.active
            
            # Escribir datos
            for row_offset, row_data in enumerate(data):
                for col_offset, value in enumerate(row_data):
                    row = start_row + row_offset
                    col = start_col + col_offset
                    ws.cell(row=row, column=col, value=value)
                    
            wb.save(str(filepath))
            wb.close()
            
            return {
                "success": True,
                "filepath": str(filepath),
                "cells_written": sum(len(row) for row in data)
            }
            
        except ImportError:
            return {"success": False, "error": "openpyxl not installed"}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def _create_word(self, task: Dict) -> Dict:
        """Crea un documento Word nuevo."""
        filename = task.get("filename", "document.docx")
        content = task.get("content", [])
        
        try:
            from docx import Document
            
            doc = Document()
            
            # Agregar contenido
            for item in content:
                if item.get("type") == "heading":
                    level = item.get("level", 1)
                    text = item.get("text", "")
                    doc.add_heading(text, level=level)
                elif item.get("type") == "paragraph":
                    text = item.get("text", "")
                    doc.add_paragraph(text)
                elif item.get("type") == "table":
                    data = item.get("data", [])
                    if data:
                        table = doc.add_table(rows=len(data), cols=len(data[0]))
                        for i, row in enumerate(data):
                            for j, cell in enumerate(row):
                                table.cell(i, j).text = str(cell)
                                
            # Guardar
            filepath = Path(filename)
            if not filepath.is_absolute():
                filepath = Path.cwd() / filepath
                
            doc.save(str(filepath))
            
            logger.info(f"Word document created: {filepath}")
            return {
                "success": True,
                "filepath": str(filepath),
                "elements": len(content)
            }
            
        except ImportError:
            return {"success": False, "error": "python-docx not installed"}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def _read_word(self, task: Dict) -> Dict:
        """Lee un documento Word."""
        filename = task.get("filename")
        
        if not filename:
            return {"success": False, "error": "Filename required"}
            
        try:
            from docx import Document
            
            filepath = Path(filename)
            if not filepath.exists():
                return {"success": False, "error": "File not found"}
                
            doc = Document(str(filepath))
            
            content = []
            for para in doc.paragraphs:
                content.append({
                    "type": "paragraph",
                    "text": para.text
                })
                
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                content.append({
                    "type": "table",
                    "data": table_data
                })
                
            return {
                "success": True,
                "content": content,
                "paragraphs": len([c for c in content if c["type"] == "paragraph"]),
                "tables": len([c for c in content if c["type"] == "table"])
            }
            
        except ImportError:
            return {"success": False, "error": "python-docx not installed"}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def _edit_word(self, task: Dict) -> Dict:
        """Edita un documento Word existente."""
        filename = task.get("filename")
        operations = task.get("operations", [])
        
        if not filename:
            return {"success": False, "error": "Filename required"}
            
        try:
            from docx import Document
            
            filepath = Path(filename)
            if not filepath.exists():
                return {"success": False, "error": "File not found"}
                
            doc = Document(str(filepath))
            
            for op in operations:
                op_type = op.get("type")
                
                if op_type == "add_paragraph":
                    doc.add_paragraph(op.get("text", ""))
                elif op_type == "add_heading":
                    doc.add_heading(op.get("text", ""), level=op.get("level", 1))
                elif op_type == "replace_text":
                    old_text = op.get("old", "")
                    new_text = op.get("new", "")
                    for para in doc.paragraphs:
                        if old_text in para.text:
                            para.text = para.text.replace(old_text, new_text)
                            
            doc.save(str(filepath))
            
            return {
                "success": True,
                "filepath": str(filepath),
                "operations_completed": len(operations)
            }
            
        except ImportError:
            return {"success": False, "error": "python-docx not installed"}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def shutdown(self):
        """Detiene el agente limpiando recursos."""
        logger.info(f"Shutting down {self.name}...")
        self.current_document = None
        self.document_type = None
        self.state = AgentState.STOPPED
        logger.info(f"{self.name} stopped.")
